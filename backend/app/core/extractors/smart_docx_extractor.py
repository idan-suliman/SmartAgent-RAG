# backend/app/core/extractors/smart_docx_extractor.py
# -*- coding: utf-8 -*-
from __future__ import annotations

import io
import os
import zipfile
from pathlib import Path
from typing import List

import pytesseract
from PIL import Image
from docx import Document

# Import shared utilities
from backend.app.core.text_utils import clean_text_output, evaluate_quality

# -------------------------
# Configuration
# -------------------------
TESSERACT_CMD = os.getenv("TESSERACT_CMD", "").strip()
if TESSERACT_CMD:
    pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD

OCR_LANG = "heb+eng"


# -------------------------
# Extraction (DOCX)
# -------------------------
def extract_text_docx(docx_path: str) -> str:
    """Extracts text from paragraphs and tables in a .docx file."""
    try:
        doc = Document(docx_path)
        parts: List[str] = []

        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                parts.append(t)

        for table in doc.tables:
            for row in table.rows:
                # Join cell content with spacing
                row_text = "  ".join((cell.text or "").strip() for cell in row.cells if (cell.text or "").strip())
                if row_text.strip():
                    parts.append(row_text.strip())

        return "\n".join(parts).strip()
    except Exception:
        return ""


def extract_images_from_docx(docx_path: str) -> List[tuple[str, bytes]]:
    """
    Extracts embedded images from the DOCX package (which is a ZIP file).
    Target: word/media/ folder inside the zip.
    """
    images: List[tuple[str, bytes]] = []
    try:
        with zipfile.ZipFile(docx_path, "r") as z:
            for name in z.namelist():
                low = name.lower()
                # Filter for image files in the media folder
                if low.startswith("word/media/") and low.split(".")[-1] in {"png", "jpg", "jpeg"}:
                    images.append((Path(name).name, z.read(name)))
    except Exception:
        pass
    return images


def ocr_images(images: List[tuple[str, bytes]], lang: str = OCR_LANG) -> str:
    """Run OCR on extracted images and return concatenated text."""
    out: List[str] = []
    for _, data in images:
        try:
            img = Image.open(io.BytesIO(data))
            text = pytesseract.image_to_string(img, lang=lang, config="--psm 6").strip()
            if text:
                out.append(text)
        except Exception:
            continue
    return "\n".join(out).strip()


# -------------------------
# Public API
# -------------------------
def extract_text_from_docx(docx_path: str) -> str:
    """
    Main Entry Point for DOCX:
    1. Extract standard text.
    2. If text is sparse (likely a scanned document inside DOCX), extract and OCR images.
    3. Return cleaned text.
    """
    src = Path(docx_path)
    text = extract_text_docx(str(src))

    try:
        # Check quality. If fails, it might be an image-based DOCX.
        is_good, _ = evaluate_quality(text, file_size_bytes=src.stat().st_size)
        
        if not is_good:
            print(f"[DOCX] Low text quality for {src.name}, attempting image extraction...")
            images = extract_images_from_docx(str(src))
            if images:
                ocr_text = ocr_images(images)
                if ocr_text:
                    # Append OCR text to whatever (if any) text we found
                    text = (text + "\n" + ocr_text).strip()
    except Exception:
        # Fail gracefully, return whatever text we have
        pass

    return clean_text_output(text)


__all__ = ["extract_text_from_docx"]